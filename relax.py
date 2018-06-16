#!/usr/bin/env python
import sys
import nmap

from docx import Document
from docx.shared import Inches

VER = 2

try:
    if sys.version_info >= (3, 0):
        VER = 3
        from urllib.request import urlopen
        from urllib.error import URLError
        raw_input = input
    else:
        from urllib2 import urlopen
        from urllib2 import URLError
except:
        pass

documento = Document()

def validarespuesta(sitios):
    count = 0
    for sitio in sitios:
        count = count + 1
        print str(count) + ': ' + sitio
    respuesta = raw_input('\n' +
                          '\033[91mEs correcta la lista de los sitios? (S/N): \033[0m')
    if respuesta.lower() == 's':
        comenzar_escaneo(sitios)
    elif respuesta.lower() == 'n':
        index = raw_input('\033[91mQue sitio esta'
                          ' mal? Introduce el numero identificador:\033[0m ')
        sitios.pop(int(index)-1)
        validarespuesta(sitios)
    else:
        print 'Escriba un comando correcto por favor'
        validarespuesta(sitios)

def fetch(url, decoding='utf-8'):
    "Obtiene el contenido de la url"
    return urlopen(url).read().decode(decoding)

def scannmap(host):
    "Realiza el escaneo del host con nmap"
    print '\033[92m[+]  \033[0mComenzando escaneo con Nmap de: ' + host + '\033[92m[+]\033[0m'
    scan = nmap.PortScanner()
    scan.scan(host)
    for host in scan.all_hosts():
        print '----------------------------------------------------'
        print 'Host : %s (%s)' % (host, scan[host].hostname())
        print 'Estado : %s' % scan[host].state()
        for proto in scan[host].all_protocols():
            print '----------'
            print 'Protocolo : %s' % proto
            lport = scan[host][proto].keys()
            lport.sort()
            for port in lport:
                print 'Puerto : %s\tEstado : %s\tUtilizado para: %state - %s' % (port, scan[host][proto][port]['state'], scan[host][proto][port]['product'], scan[host][proto][port]['version'])
                documento.add_paragraph('Puerto : %s\tEstado : %s\tUtilizado para: %state - %s' % (port, scan[host][proto][port]['state'], scan[host][proto][port]['product'], scan[host][proto][port]['version']))

def whois(host):
    "Realiza un whois al host"
    print '\n\033[92m[+]\033[0mComenzando escaneo con WHOIS de: ' + host + '\033[92m[+] \033[0m'
    url = "http://api.hackertarget.com/whois/?q=" + host
    pwho = fetch(url)
    documento.add_paragraph(pwho)

def dnslookup(host):
    print '\n\033[92m[+]\033[0mComenzando escaneo con DNSLOOKUP de: ' + host + '\033[92m[+] \033[0m'
    "Realiza un dnslookup al host"
    ns = "http://api.hackertarget.com/dnslookup/?q=" + host
    pns = fetch(ns)
    documento.add_paragraph(pns)

def page_links(host):
    print '\n\033[92m[+]\033[0mComenzando a escanear las URL de la pagina: ' + host + '\033[92m[+] \033[0m'
    ns = "https://api.hackertarget.com/pagelinks/?q=" + host
    res = fetch(ns)
    documento.add_paragraph(res)

def test_ping(host):
    print '\n\033[92m[+]\033[0mComenzando el testeo de los ping de: ' + host + '\033[92m[+] \033[0m'
    ns = "https://api.hackertarget.com/nping/?q=" + host
    res = fetch(ns)
    documento.add_paragraph(res)

def tracerouter(host):
    print '\n\033[92m[+]\033[0mRealizando tracer route sobre: ' + host + '\033[92m[+] \033[0m'
    ns = "https://api.hackertarget.com/mtr/?q=" + host
    res = fetch(ns)
    documento.add_paragraph(res)


def comenzar_escaneo(sitios):
    "Comienza el escaneo completo de los sitios ingresados"
    for host in sitios:
        documento.add_heading("Resultados del analisis con NMAP")
        scannmap(host)
        documento.add_heading("Resultados del analisis con WHOIS")
        whois(host)
        documento.add_heading("Resultados del analisis con DNSLOOKUP")
        dnslookup(host)
        documento.add_heading("Resultados del analisis con TRACEROUTER")
        tracerouter(host)
        documento.add_heading("Resultados del analisis con TEST PING")
        test_ping(host)
        documento.add_heading("Resultados del analisis con PAGE LINKS")
        page_links(host)
        documento.save('Analisis - '+ host + '.docx')

def banner():
    print '\033[92m/$$$$       /$$$$\033[0m       /$$$$$$$            | $$                          \033[92m /$$$$       /$$$$\033[0m'
    print '\033[92m| $$_/  /$$ |_  $$\033[0m      | $$__  $$          | $$                          \033[92m| $$_/  /$$ |_  $$\033[0m'
    print '\033[92m| $$   | $$   | $$\033[0m      | $$  \ $$  /$$$$$$ | $$  /$$$$$$  /$$   /$$      \033[92m| $$   | $$   | $$\033[0m'
    print '\033[92m| $$ /$$$$$$$$| $$\033[0m      | $$$$$$$/ /$$__  $$| $$ |____  $$|  $$ /$$/      \033[92m| $$ /$$$$$$$$| $$\033[0m'
    print '\033[92m| $$|__  $$__/| $$\033[0m      | $$__  $$| $$$$$$$$| $$  /$$$$$$$ \  $$$$/       \033[92m| $$|__  $$__/| $$\033[0m'
    print '\033[92m| $$   | $$   | $$\033[0m      | $$  \ $$| $$_____/| $$ /$$__  $$  >$$  $$       \033[92m| $$   | $$   | $$\033[0m'
    print '\033[92m| $$$$ |__/  /$$$$\033[0m      | $$  | $$|  $$$$$$$| $$|  $$$$$$$ /$$/\  $$      \033[92m| $$$$ |__/  /$$$$\033[0m'
    print '\033[92m|____/      |____/\033[0m      |__/  |__/ \_______/|__/ \_______/|__/  \__/      \033[92m|____/      |____/\033[0m'
    print '\t\t\t\tHecho con mucho \033[91m<3\033[0m por Hector Camacho!'
    print '\n'
    menu()

def inicio():
    banner()
    sitios = []
    sitio = raw_input('\033[91mIntroduce algun sitio '
                      'para analizar:\033[0m ')
    sitios.append(sitio)
    while sitio != 'Listo' and sitio != 'Salir':
        sitio = raw_input('\033[91mIntroduce algun '
                          'sitio para analizar:\033[0m ')
        if sitio == 'Listo':
            validarespuesta(sitios)
            break;
        if sitio == 'Salir':
            return 'Saliendo...';
            break;
        else:
            sitios.append(sitio)

def menu():
    print '\n'
    print 'Una vez que los sitios esten capturados escribe "Listo" para continuar con el escaneo.\n'
    print 'Para salir del programa escribe "Salir"\n'

if __name__ == '__main__':
    inicio()
